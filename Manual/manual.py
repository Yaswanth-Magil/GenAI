import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
import numpy as np
import os
from docx.enum.table import WD_ALIGN_VERTICAL
# Define the file paths for the three months of data
file_paths = ["/Users/yash/Downloads/Today/Back Then/A2b December Month.xlsx", "/Users/yash/Downloads/Today/Back Then/January/A2b January month.xlsx", "/Users/yash/Downloads/Today/Splitted/GenAI/A2b Feb month.xlsx"]
# Initialize lists to store counts and percentages
months = ['December', 'January', 'February']
positive_counts = []
negative_counts = []
neutral_counts = []

positive_percent = []
negative_percent = []
neutral_percent = []

month_denoted = months[2]

sheets = ['Princeton', 'SPF', 'Chicago', 'Parsippany']
sheet = sheets[0]

# Function to calculate sentiment counts and percentages
def calculate_sentiments(df):
    positive = df[df['Sentiment'] == 'Positive'].shape[0]
    negative = df[df['Sentiment'] == 'Negative'].shape[0]
    neutral = df[df['Sentiment'] == 'Neutral'].shape[0]
    total = positive + negative + neutral
    
    positive_counts.append(positive)
    negative_counts.append(negative)
    neutral_counts.append(neutral)
    
    positive_percent.append(positive / total * 100)
    negative_percent.append(negative / total * 100)
    neutral_percent.append(neutral / total * 100)

# Read the data from the Excel files and calculate sentiments
for file_path in file_paths:
    df = pd.read_excel(file_path, sheet_name=sheet)
    calculate_sentiments(df)

# Create a Word document
doc = Document()

# Access the first section
section = doc.sections[0]

section = doc.sections[0]
section.top_margin = Cm(0.5)
section.bottom_margin = Cm(0.5)
section.header_distance = Cm(0.5)
section.footer_distance = Cm(0.5)

# Add header image
    # Access the first section
   
    
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

header_img_path = '/Users/yash/Downloads/Today/Back Then/maghilLogo.png'  # Header image
footer_img_path =  '/Users/yash/Downloads/Today/Back Then/og.png'
header_img = Image.open(header_img_path)
header_img_byte_arr = io.BytesIO()
header_img.save(header_img_byte_arr, format='PNG')
header_img_byte_arr.seek(0)


table = header.add_table(rows=1, cols=3, width=Inches(4.33))
table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
table.cell(0,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

cell1 = table.cell(0,0)
cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p1 = cell1.paragraphs[0]
run1 = p1.add_run()
run1.add_picture(header_img_byte_arr, width=Inches(1.01), height=Inches(0.5))

cell2 = table.cell(0,1)
cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell2.text = "                                                                     "
# Add footer image
footer_img = Image.open(footer_img_path)
footer_img_byte_arr = io.BytesIO()
footer_img.save(footer_img_byte_arr, format='PNG')
footer_img_byte_arr.seek(0)

cell3 = table.cell(0,2)
cell3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p2 = cell3.paragraphs[0]
run2 = p2.add_run()
run2.add_picture(footer_img_byte_arr, width=Inches(1.46), height=Inches(0.4))


table.columns[0].width = Inches(3)
table.columns[1].width = Inches(3)

# Remove Footer image
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = ""

# Add title and center it
title = doc.add_heading('A2B'+" "+sheet+" - "+'GenAI Insights', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

subtitle2 = doc.add_heading('Reporting Month:'+month_denoted+' 2024', level=2)
for run in subtitle2.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph('Reviews Considered')

# Add table for review sources
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Format table header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Source'
hdr_cells[1].text = 'Count'

# Format header text
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add table data
source_counts = df['Source'].value_counts()

google_count = source_counts.get('Google', 0)
grubhub_count = source_counts.get('Grughub', 0)
yelp_count = source_counts.get('Yelp', 0)
uber_count = source_counts.get('Uber', 0)
doordash_count = source_counts.get('Doordash', 0)

# Calculate the total count of all sources
total_count = google_count + grubhub_count + yelp_count + uber_count + doordash_count

# Filter out sources with count 0
data = [
    ['Google', google_count],
    ['Yelp', yelp_count],
    ['Grubhub', grubhub_count],
    ['Uber', uber_count],
    ['Doordash', doordash_count]
]

filtered_data = [item for item in data if item[1] > 0]

# Add the total count row separately
filtered_data.append(['Total', total_count])

for source, count in filtered_data:
    row_cells = table.add_row().cells
    row_cells[0].text = source
    row_cells[1].text = str(count)
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if source == 'Total':  # Check if the row is 'Total'
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True 

# Add extra space after the table
doc.add_paragraph('')

# Add headings with black font color
heading = doc.add_heading('Reviews Sentimental Analysis Summary', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

heading = doc.add_heading('     Review Sentiments Over Three Months                                Sentiment Distribution of '+months[0], level=3)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

# Plotting the charts
fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(14, 6))  # Create two subplots

# Plotting the stacked bar chart
bar_width = 0.5
index = np.arange(len(months))

bar_negative = ax1.bar(index, negative_percent, bar_width, label='Negative', color='red')
bar_neutral = ax1.bar(index, neutral_percent, bar_width, bottom=negative_percent, label='Neutral', color='yellow')
bar_positive = ax1.bar(index, positive_percent, bar_width, bottom=np.array(negative_percent) + np.array(neutral_percent), label='Positive', color='green')

# Adding annotations for counts and percentages
for i, month in enumerate(months):
    negative_height = negative_percent[i]
    neutral_height = neutral_percent[i]
    positive_height = positive_percent[i]
    
    negative_center = negative_height / 2
    neutral_center = negative_height + neutral_height / 2
    positive_center = negative_height + neutral_height + positive_height / 2

    ax1.text(i, negative_center, f'{negative_counts[i]} ({negative_percent[i]:.1f}%)', ha='center', color='white')
    ax1.text(i, neutral_center, f'{neutral_counts[i]} ({neutral_percent[i]:.1f}%)', ha='center', color='black')
    ax1.text(i, positive_center, f'{positive_counts[i]} ({positive_percent[i]:.1f}%)', ha='center', color='black')

ax1.set_xlabel('Month')
ax1.set_ylabel('Percentage of Reviews')
ax1.set_ylim(0, 100)
ax1.set_xticks(index)
ax1.set_xticklabels(months)
ax1.legend()
ax1.grid(False)

# Plotting the pie chart
labels = ['Positive', 'Negative', 'Neutral']
sizes = [positive_percent[-1], negative_percent[-1], neutral_percent[-1]]
colors = ['green', 'red', 'yellow']
explode = (0, 0.2, 0.1)  # explode the 2nd slice (i.e. 'Negative') more than the 3rd slice ('Neutral')

ax2.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
ax2.text(-2.0, -1.5, f'Positive: {positive_counts[-1]} ({positive_percent[-1]:.1f}%)\nNegative: {negative_counts[-1]} ({negative_percent[-1]:.1f}%)\nNeutral: {neutral_counts[-1]} ({neutral_percent[-1]:.1f}%)', fontsize=12, bbox=dict(facecolor='white', alpha=0.5), color='black')

plt.tight_layout()
chart_img_path = '/Users/yash/MagilHub/Project 2/GenAI/GenAIchart.png'
plt.savefig(chart_img_path)
plt.close()

# Add the chart image to the Word document
doc.add_picture(chart_img_path, width=Inches(6.0))

# Adding space after the image
doc.add_paragraph('')

heading = doc.add_heading('Most Common Dishes Mentioned', level=3)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

attributes_1 = [
    "Dosai", 
    "Idli",      # 1
    "Mini Tiffin",   # 2
    "Adai Avial",
    "navratan korma",

]



counts_positive_1 = [5,5,3,0,0,]
counts_negative_1 = [0,0, 0,1,1,]
counts_neutral_1 = [0,0,0, 0, 0]


attributes_2 = [
    "Ali",       
    "Kiyas",   
    "Rizmi",
    "Hakeem",    
]

counts_positive_2 = [210, 24, 7, 5]
counts_negative_2 = [0, 0, 0, 1]
counts_neutral_2 = [0, 0, 0, 0]


                # Example counts of neutral mentions
# Create the combined horizontal stacked bar chart
fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(14, 6))

# Plot the first horizontal stacked bar chart
ax1.barh(attributes_1, counts_positive_1, color='green', label='Positive')
ax1.barh(attributes_1, counts_neutral_1, left=counts_positive_1, color='yellow', label='Neutral')
ax1.barh(attributes_1, counts_negative_1, left=np.array(counts_positive_1) + np.array(counts_neutral_1), color='red', label='Negative')

# Adding annotations for counts
for i, attribute in enumerate(attributes_1):
    ax1.text(counts_positive_1[i] / 2, i, str(counts_positive_1[i]), ha='center', va='center', color='white', fontsize=10)
    ax1.text(counts_positive_1[i] + counts_neutral_1[i] / 2, i, str(counts_neutral_1[i]), ha='center', va='center', color='black', fontsize=10)
    ax1.text(counts_positive_1[i] + counts_neutral_1[i] + counts_negative_1[i] / 2, i, str(counts_negative_1[i]), ha='center', va='center', color='white', fontsize=10)

ax1.set_xlabel('Count')
ax1.set_ylabel('Foods')
ax1.legend()
ax1.grid(False)

# Plot the second horizontal stacked bar chart
ax2.barh(attributes_2, counts_positive_2, color='green', label='Positive')
ax2.barh(attributes_2, counts_neutral_2, left=counts_positive_2, color='yellow', label='Neutral')
ax2.barh(attributes_2, counts_negative_2, left=np.array(counts_positive_2) + np.array(counts_neutral_2), color='red', label='Negative')

# Adding annotations for counts
for i, attribute in enumerate(attributes_2):
    ax2.text(counts_positive_2[i] / 2, i, str(counts_positive_2[i]), ha='center', va='center', color='white', fontsize=10)
    ax2.text(counts_positive_2[i] + counts_neutral_2[i] / 2, i, str(counts_neutral_2[i]), ha='center', va='center', color='black', fontsize=10)
    ax2.text(counts_positive_2[i] + counts_neutral_2[i] + counts_negative_2[i] / 2, i, str(counts_negative_2[i]), ha='center', va='center', color='white', fontsize=10)

ax2.set_xlabel('Count')
ax2.set_ylabel('Staffs')
ax2.legend()
ax2.grid(False)

plt.tight_layout()

# Save the combined horizontal stacked bar chart
horizontal_combined_chart_img_path = '/Users/yash/MagilHub/Project 2/GenAI/GenAI/horizontal_stacked_bar.png'
plt.savefig(horizontal_combined_chart_img_path)
plt.close()

# Add the combined horizontal stacked bar chart image to the Word document
doc.add_picture(horizontal_combined_chart_img_path, width=Inches(6.0))

# Add a page break to start the second page
doc.add_page_break()# Data for the vertical stacked bar chart
# days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
# positive_counts_days = [3, 13, 3, 101, 10, 1, 1]
# negative_counts_days = [0, 2, 0, 12, 0, 1, 0]
# neutral_counts_days = [0, 0, 1, 4, 0, 0, 0]

# # Create the vertical stacked bar chart
# fig, ax = plt.subplots(figsize=(10, 6))

# bar_width = 0.5
# index = np.arange(len(days))

# bar_positive_days = ax.bar(index, positive_counts_days, bar_width, label='Positive', color='green')
# bar_neutral_days = ax.bar(index, neutral_counts_days, bar_width, bottom=positive_counts_days, label='Neutral', color='yellow')
# bar_negative_days = ax.bar(index, negative_counts_days, bar_width, bottom=np.array(positive_counts_days) + np.array(neutral_counts_days), label='Negative', color='red')

# # Adding annotations for counts
# for i, day in enumerate(days):
#     positive_height = positive_counts_days[i]
#     neutral_height = neutral_counts_days[i]
#     negative_height = negative_counts_days[i]
    
#     positive_center = positive_height / 2
#     neutral_center = positive_height + neutral_height / 2
#     negative_center = positive_height + neutral_height + negative_height / 2

#     ax.text(i, positive_center, str(positive_counts_days[i]), ha='center', va='center', color='white')
#     ax.text(i, neutral_center, str(neutral_counts_days[i]), ha='center', va='center', color='black')
#     ax.text(i, negative_center, str(negative_counts_days[i]), ha='center', va='center', color='white')

# ax.set_xlabel('Days of the Week')
# ax.set_ylabel('Count')
# ax.set_xticks(index)
# ax.set_xticklabels(days)
# ax.legend()
# ax.grid(False)

# plt.tight_layout()

# # Save the vertical stacked bar chart
# vertical_stacked_chart_img_path = 'C:/Users/shree/Downloads/vertical_stacked_chart.png'
# plt.savefig(vertical_stacked_chart_img_path)
# plt.close()

# # Add the vertical stacked bar chart image to the second page of the Word document
# doc.add_picture(vertical_stacked_chart_img_path, width=Inches(6.0))
#############################################################################################################################
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches


categories = [
    "Food Quality",           # 1 negative mention
    "Service",
    "Speed of Service"              # 1 negative mention
    "Staff Friendliness",           # 1 negative mention
    "Overall Experience",           # 1 negative mention
    "Ambiance",     # 1 negative mention
                   # 2 negative mentions
               # 3 negative mentions
]

positive_counts = [199, 190, 20,  121, 90, 20,]
negative_counts = [3,3,3, 1,2,2,]
   # Adjusted based on observed negative feedback


# Plot
fig, ax = plt.subplots(figsize=(12, 6))

# Positions of bars on the y-axis
y_pos = np.arange(len(categories))

# Create horizontal bars with positive and negative counts on the same line
bar_width = 0.4
bars1 = ax.barh(y_pos, positive_counts, bar_width, color='green', edgecolor='black', label='Positive')
bars2 = ax.barh(y_pos, -np.array(negative_counts), bar_width, color='red', edgecolor='black', label='Negative')

# Add a line to separate positive and negative
ax.axvline(0, color='black', linewidth=0.8)

# Add labels
ax.set_yticks(y_pos)
ax.set_yticklabels(categories)
ax.set_xlabel('Counts')
ax.set_title('Positive and Negative Mentions by Category')

# Add counts on the bars if not 0
for bar in bars1:
    width = bar.get_width()
    if width != 0:
        ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, f'{width}', va='center', ha='left')

for bar in bars2:
    width = bar.get_width()
    if width != 0:
        ax.text(width - 0.1, bar.get_y() + bar.get_height()/2, f'{-width}', va='center', ha='right')

# Add a legend
ax.legend()

# Adjust layout to make sure y-axis labels are fully visible
plt.tight_layout()

# Save the plot to a PNG file
plot_filename = '/Users/yash/MagilHub/Project 2/GenAI/GenAI/plot.png'
plt.savefig(plot_filename)

# Save the plot image to a docx file
doc.add_paragraph()
doc.add_picture(plot_filename, width=Inches(6), height=Inches(5))


# Save the Word document
#A2B_Princeton_October.docx
#A2B_Princeton_Justification.docx
output_path = '/Users/yash/Downloads/Today/Splitted/GenAI/Manual/Outputs/A2B_Princeton_February.docx'
doc.save(output_path)

print('Document created and saved at:', output_path)