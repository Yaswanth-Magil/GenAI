import mysql.connector
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from PIL import Image
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import matplotlib.pyplot as plt
import numpy as np
import json
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import subprocess
import logging
import re
#from docx.enum.text import WD_BULLET_TYPE #Removed this as this import gives error
#from docx.enum.text import WD_BREAK_TYPE

# Database credentials - Replace with your actual credentials
# db_host = "localhost"
# db_user = "root"
# db_password = "Yaswanth123."
# db_name = "genai"
# table_name = "output_dummy_2"

db_host = "10.162.0.3"
db_user = "root"
db_password = "Z*ZlRmnFCP@9V"
db_name = "mhrq"
table_name = "output_genai"

year = datetime.datetime.now().year

# Header Image Paths - Replace with your actual image paths
header_img_path = '/Users/yash/Downloads/Today/Back Then/maghilLogo.png'
footer_img_path = '/Users/yash/Downloads/Today/Back Then/og.png'  # Footer image

# Columns to fetch
columns_to_fetch = [
    'positive_summary',
    'negative_summary',
    'where_i_do_better',
    'where_competitor_do_better',
    'trend_pos_to_neg',
    'trend_neg_to_pos',
    'overall_positive_count',
    'overall_negative_count',
    'overall_neutral_count',
    'dish_positive_counts',
    'dish_negative_counts',
    'staff_positive_counts',
    'staff_negative_counts',
    'category_positive_counts',
    'category_negative_counts',
    'trend_note',
    'category_note'
]


def process_text(document, text):
    """Processes text for bold patterns and bullet points."""

    # Split the text into paragraphs
    paragraphs = text.splitlines()

    for paragraph_text in paragraphs:
        if paragraph_text.strip() == "":
            continue  # Skip empty paragraphs

        # Check for bullet points
        bullet_match = re.match(r"^\*\s+(.+)", paragraph_text)  # Match "* text"
        bullet_match2 = re.match(r"^\s*\*\s+(.+)", paragraph_text)  # Match " * text" with spaces

        if bullet_match or bullet_match2:
            bullet_text = bullet_match.group(1) if bullet_match else bullet_match2.group(1)
            p = document.add_paragraph("", style='List Bullet')

            # Process bold formatting within bullet points
            parts = re.split(r"(\*\*[^*]+\*\*)", bullet_text)
            for part in parts:
                run = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
                if part.startswith("**") and part.endswith("**"):
                    run.bold = True

        else:
            p = document.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", paragraph_text)

            for part in parts:
                run = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
                if part.startswith("**") and part.endswith("**"):
                    run.bold = True


def fetch_data_from_db(outlet, review_month, num_months=3):
    """Fetches data from the database for a specified number of months."""
    data = []
    try:
        mydb = mysql.connector.connect(
            host=db_host,
            user=db_user,
            password=db_password,
            database=db_name
        )
        mycursor = mydb.cursor()

        # Fetch data for the specified number of months (up to 3)
        for i in range(num_months):
            current_month = review_month - i
            
            # Handle cases where current_month goes below 1 (December of previous year)
            if current_month < 1:
                current_month = current_month + 12 # rollover to previous year december

            # Construct the SQL query
            column_names = ", ".join(columns_to_fetch) # Join column names into a string
            sql = f"SELECT {column_names} FROM {table_name} WHERE outlet = %s AND review_month = %s"
            val = (outlet, current_month)
            mycursor.execute(sql, val)
            result = mycursor.fetchone()

            if result:
                # Convert the tuple to a dictionary for easier access by column name.
                month_data = dict(zip(columns_to_fetch, result))
                data.append((current_month, month_data)) # Append month and data

        # Sort the data by month in ascending order
        data.sort(key=lambda x: x[0])

        return data

    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return None
    finally:
        if mydb and mydb.is_connected():  # check if mydb is defined before calling is_connected()
            if mycursor:  # check if mycursor is defined before closing it
                mycursor.close()
            mydb.close()


def create_trend_chart(data, chart_img_path):
    """Generates a trend chart using the fetched data for multiple months."""
    if not data or len(data) == 0:
        return False

    months = [str(month) for month, _ in data] # get the months from data.
    positive_counts = [month_data['overall_positive_count'] or 0 for _, month_data in data]
    negative_counts = [month_data['overall_negative_count'] or 0 for _, month_data in data]
    neutral_counts = [month_data['overall_neutral_count'] or 0 for _, month_data in data]

    total_counts = [p + n + neu for p, n, neu in zip(positive_counts, negative_counts, neutral_counts)]
    positive_percent = [p / total * 100 if total else 0 for p, total in zip(positive_counts, total_counts)]
    negative_percent = [n / total * 100 if total else 0 for n, total in zip(negative_counts, total_counts)]
    neutral_percent = [neu / total * 100 if total else 0 for neu, total in zip(neutral_counts, total_counts)]

    # Plotting the charts
    fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(14, 6))  # Create two subplots

    # Plotting the stacked bar chart
    bar_width = 0.5
    index = np.arange(len(months))

    bar_negative = ax1.bar(index, negative_percent, bar_width, label='Negative', color='red')
    bar_neutral = ax1.bar(index, neutral_percent, bar_width, bottom=negative_percent, label='Neutral',
                           color='yellow')
    bar_positive = ax1.bar(index, positive_percent, bar_width,
                           bottom=np.array(negative_percent) + np.array(neutral_percent), label='Positive',
                           color='green')

    # Adding annotations for counts and percentages
    for i, month in enumerate(months):
        negative_height = negative_percent[i]
        neutral_height = neutral_percent[i]
        positive_height = positive_percent[i]

        negative_center = negative_height / 2
        neutral_center = negative_height + neutral_height / 2
        positive_center = negative_height + neutral_height + positive_height / 2

        ax1.text(i, negative_center, f'{negative_counts[i]} ({negative_percent[i]:.1f}%)', ha='center',
                 color='white')
        ax1.text(i, neutral_center, f'{neutral_counts[i]} ({neutral_percent[i]:.1f}%)', ha='center', color='black')
        ax1.text(i, positive_center, f'{positive_counts[i]} ({positive_percent[i]:.1f}%)', ha='center', color='black')

    ax1.set_xlabel('Month')
    ax1.set_ylabel('Percentage of Reviews')
    ax1.set_ylim(0, 100)
    ax1.set_xticks(index)
    ax1.set_xticklabels(months)
    ax1.legend()
    ax1.grid(False)

    # Plotting the pie chart (for the latest month)
    labels = ['Positive', 'Negative', 'Neutral']
    sizes = [positive_percent[-1], negative_percent[-1], neutral_percent[-1]]
    colors = ['green', 'red', 'yellow']
    explode = (0, 0.2, 0.1)  # explode the 2nd slice (i.e. 'Negative') more than the 3rd slice ('Neutral')

    ax2.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    ax2.text(-2.0, -1.5,
             f'Positive: {positive_counts[-1]} ({positive_percent[-1]:.1f}%)\nNegative: {negative_counts[-1]} ({negative_percent[-1]:.1f}%)\nNeutral: {neutral_counts[-1]} ({neutral_percent[-1]:.1f}%)',
             fontsize=12, bbox=dict(facecolor='white', alpha=0.5), color='black')

    plt.tight_layout()
    plt.savefig(chart_img_path)
    plt.close()
    return True


def create_most_mentioned_chart(data, horizontal_combined_chart_img_path):
    """Generates a chart of most mentioned staff and dishes."""
    if not data:
        return False

    # Extract data from JSON columns, handling potential None values and empty JSON
    # For multi-month analysis, you might want to aggregate the counts across all months
    # For simplicity, we'll use data from the latest month. You can modify as needed.
    latest_month_data = data[-1][1]  # Access data of the latest month
    dish_positive_counts = json.loads(latest_month_data['dish_positive_counts']) if latest_month_data['dish_positive_counts'] else {}
    dish_negative_counts = json.loads(latest_month_data['dish_negative_counts']) if latest_month_data['dish_negative_counts'] else {}
    staff_positive_counts = json.loads(latest_month_data['staff_positive_counts']) if latest_month_data['staff_positive_counts'] else {}
    staff_negative_counts = json.loads(latest_month_data['staff_negative_counts']) if latest_month_data['staff_negative_counts'] else {}

    # Process dishes (Top 5)
    dish_counts = {k: dish_positive_counts.get(k, 0) + dish_negative_counts.get(k, 0) for k in
                   set(dish_positive_counts) | set(dish_negative_counts)}
    top_dishes = sorted(dish_counts, key=dish_counts.get, reverse=True)[:5]  # Get top 5
    attributes_1 = top_dishes
    counts_positive_1 = [dish_positive_counts.get(attr, 0) for attr in attributes_1]
    counts_negative_1 = [dish_negative_counts.get(attr, 0) for attr in attributes_1]
    counts_neutral_1 = [0] * len(attributes_1)  # Assuming no neutral counts in JSON

    # Process staff (Top 3)
    staff_counts = {k: staff_positive_counts.get(k, 0) + staff_negative_counts.get(k, 0) for k in
                    set(staff_positive_counts) | set(staff_negative_counts)}
    top_staff = sorted(staff_counts, key=staff_counts.get, reverse=True)[:3]  # Get top 3
    attributes_2 = top_staff
    counts_positive_2 = [staff_positive_counts.get(attr, 0) for attr in attributes_2]
    counts_negative_2 = [staff_negative_counts.get(attr, 0) for attr in attributes_2]
    counts_neutral_2 = [0] * len(attributes_2)  # Assuming no neutral counts in JSON

    # Create the combined horizontal stacked bar chart
    fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(14, 6))

    # Plot the first horizontal stacked bar chart
    ax1.barh(attributes_1, counts_positive_1, color='green', label='Positive')
    ax1.barh(attributes_1, counts_neutral_1, left=counts_positive_1, color='yellow', label='Neutral')
    ax1.barh(attributes_1, counts_negative_1,
             left=np.array(counts_positive_1) + np.array(counts_neutral_1), color='red', label='Negative')

    # Adding annotations for counts
    for i, attribute in enumerate(attributes_1):
        ax1.text(counts_positive_1[i] / 2, i, str(counts_positive_1[i]), ha='center', va='center', color='white',
                 fontsize=10)
        ax1.text(counts_positive_1[i] + counts_neutral_1[i] / 2, i, str(counts_neutral_1[i]), ha='center',
                 va='center', color='black', fontsize=10)
        ax1.text(counts_positive_1[i] + counts_neutral_1[i] + counts_negative_1[i] / 2, i,
                 str(counts_negative_1[i]), ha='center', va='center', color='white', fontsize=10)

    ax1.set_xlabel('Count')
    ax1.set_ylabel('Foods')
    ax1.legend()
    ax1.grid(False)

    # Plot the second horizontal stacked bar chart
    ax2.barh(attributes_2, counts_positive_2, color='green', label='Positive')
    ax2.barh(attributes_2, counts_neutral_2, left=counts_positive_2, color='yellow', label='Neutral')
    ax2.barh(attributes_2, counts_negative_2,
             left=np.array(counts_positive_2) + np.array(counts_neutral_2), color='red', label='Negative')

    # Adding annotations for counts
    for i, attribute in enumerate(attributes_2):
        ax2.text(counts_positive_2[i] / 2, i, str(counts_positive_2[i]), ha='center', va='center', color='white',
                 fontsize=10)
        ax2.text(counts_positive_2[i] + counts_neutral_2[i] / 2, i, str(counts_neutral_2[i]), ha='center',
                 va='center', color='black', fontsize=10)
        ax2.text(counts_positive_2[i] + counts_neutral_2[i] + counts_negative_2[i] / 2, i,
                 str(counts_negative_2[i]), ha='center', va='center', color='white', fontsize=10)

    ax2.set_xlabel('Count')
    ax2.set_ylabel('Staffs')
    ax2.legend()
    ax2.grid(False)

    plt.tight_layout()

    # Save the combined horizontal stacked bar chart
    plt.savefig(horizontal_combined_chart_img_path)
    plt.close()
    return True


def create_category_chart(data, plot_filename):
    """Generates a chart for categories mentioned, sorted by negative reviews."""
    if not data:
        return False

    # For multi-month analysis, you might want to aggregate the counts across all months
    # For simplicity, we'll use data from the latest month. You can modify as needed.
    latest_month_data = data[-1][1]  # Access data of the latest month
    # Extract data from JSON, handling potential None values
    category_positive_counts = json.loads(latest_month_data['category_positive_counts']) if latest_month_data['category_positive_counts'] else {}
    category_negative_counts = json.loads(latest_month_data['category_negative_counts']) if latest_month_data['category_negative_counts'] else {}

    # Combine positive and negative counts and sort by negative counts in descending order
    categories = list(category_positive_counts.keys())
    sorted_categories = sorted(categories, key=lambda cat: category_negative_counts.get(cat, 0),)

    positive_counts = [category_positive_counts.get(cat, 0) for cat in sorted_categories]
    negative_counts = [category_negative_counts.get(cat, 0) for cat in sorted_categories]

    # Plot
    fig, ax = plt.subplots(figsize=(12, 6))

    # Positions of bars on the y-axis
    y_pos = np.arange(len(sorted_categories))

    # Create horizontal bars with positive and negative counts on the same line
    bar_width = 0.4
    bars1 = ax.barh(y_pos, positive_counts, bar_width, color='green', edgecolor='black', label='Positive')
    bars2 = ax.barh(y_pos, -np.array(negative_counts), bar_width, color='red', edgecolor='black', label='Negative')

    # Add a line to separate positive and negative
    ax.axvline(0, color='black', linewidth=0.8)

    # Add labels
    ax.set_yticks(y_pos)
    ax.set_yticklabels(sorted_categories)
    ax.set_xlabel('Counts')
    ax.set_title('Positive and Negative Mentions by Category')

    # Add counts on the bars if not 0
    for bar in bars1:
        width = bar.get_width()
        if width != 0:
            ax.text(width + 0.1, bar.get_y() + bar.get_height() / 2, f'{width}', va='center', ha='left')

    for bar in bars2:
        width = bar.get_width()
        if width != 0:
            ax.text(width - 0.1, bar.get_y() + bar.get_height() / 2, f'{-width}', va='center', ha='right')

    # Add a legend
    ax.legend()

    # Adjust layout to make sure y-axis labels are fully visible
    plt.tight_layout()

    # Save the plot to a PNG file
    plt.savefig(plot_filename)
    plt.close()
    return True


def create_word_document(outlet, review_month, data, output_filename="output.docx"):
    """Creates a Word document with the fetched data and charts in the specified order."""

    document = Document()

    # Apply Times New Roman font to the entire document
    style = document.styles.add_style('Times New Roman', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal'].font.size = Pt(12)

    for style in document.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- HEADER SECTION ---
    section = document.sections[0]  # Access the first section

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    header_img = Image.open(header_img_path)
    header_img_byte_arr = io.BytesIO()
    header_img.save(header_img_byte_arr, format='PNG')
    header_img_byte_arr.seek(0)

    table = header.add_table(rows=1, cols=3, width=Inches(4.33))
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell1 = table.cell(0, 0)
    cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run()
    run1.add_picture(header_img_byte_arr, width=Inches(1.01), height=Inches(0.5))

    cell2 = table.cell(0, 1)
    cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell2.text = "                                                                  "
    footer_img = Image.open(footer_img_path)
    footer_img_byte_arr = io.BytesIO()
    footer_img.save(footer_img_byte_arr, format='PNG')
    footer_img_byte_arr.seek(0)

    cell3 = table.cell(0, 2)
    cell3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2 = cell3.paragraphs[0]
    run2 = p2.add_run()
    run2.add_picture(footer_img_byte_arr, width=Inches(1.46), height=Inches(0.4))

    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""

    # --- DOCUMENT BODY ---
    title = document.add_heading(f"A2B {outlet} GenAI Insights", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    subhead = document.add_heading(f"Reporting Month: {review_month} {year}", level=1)  # Add Year here
    subhead.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subhead.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    # 1. --- Trend Chart, Most Mentioned Dishes and Staff, Most Mentioned Categories ---
    # --- Trend Chart ---
    chart_img_path = f'{outlet}_{review_month}_trend_chart.png'
    if create_trend_chart(data, chart_img_path):
        trend_chart_heading = document.add_heading('Sentiment Trend', level=2)
        for run in trend_chart_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
        document.add_picture(chart_img_path, width=Inches(6), height=Inches(5))

        # Add Trend Note
        trend_note = data[-1][1].get('trend_note')  # Access 'trend_note' from the latest month
        if trend_note:
            p = document.add_paragraph()
            run = p.add_run(f"Trend Note: {trend_note}")
            run.bold = True

        else:
            document.add_paragraph("No trend note available.")

    else:
        document.add_paragraph("Could not generate sentiment trend chart.")

    # --- Most Mentioned Chart ---
    horizontal_combined_chart_img_path = f'{outlet}_{review_month}_mentioned_chart.png'
    if create_most_mentioned_chart(data, horizontal_combined_chart_img_path):
        mentioned_chart_heading = document.add_heading('Most Common Dishes and Staff Mentioned', level=2)
        for run in mentioned_chart_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
        document.add_picture(horizontal_combined_chart_img_path, width=Inches(6), height=Inches(5))
    else:
        document.add_paragraph("Could not generate most mentioned chart.")

    # --- Category Chart ---
    plot_filename = f'{outlet}_{review_month}_category_chart.png'
    if create_category_chart(data, plot_filename): # Pass plot_filename
        category_chart_heading = document.add_heading('Category Analysis', level=2)
        for run in category_chart_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
        document.add_picture(plot_filename, width=Inches(6), height=Inches(5))

        #Add Category Note
        category_note = data[-1][1].get('category_note') #accessing from the latest month
        if category_note:
            p = document.add_paragraph()
            run = p.add_run(f"Category Note: {category_note}")
            run.bold = True
        
        else:
            document.add_paragraph("No category note available")
    else:
        document.add_paragraph("Could not generate category chart.")

    document.add_page_break()
    # 2. --- Review Analysis ---
    review_analysis_heading = document.add_heading('Review Analysis', level=2)
    for run in review_analysis_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    positive_summary_heading = document.add_heading('Positive Summary', level=3)
    for run in positive_summary_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can concatenate or summarize positive summaries across months
    # For simplicity, we'll use the positive summary from the latest month

    if data:
        process_text(document, str(data[-1][1]['positive_summary']) if data[-1][1]['positive_summary'] else "No data available")
    else:
        document.add_paragraph("No data available")

    negative_summary_heading = document.add_heading('Negative Summary', level=3)
    for run in negative_summary_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can concatenate or summarize negative summaries across months
    # For simplicity, we'll use the negative summary from the latest month
    if data:
        process_text(document, str(data[-1][1]['negative_summary']) if data[-1][1]['negative_summary'] else "No data available")
    else:
        document.add_paragraph("No data available")

    document.add_page_break()
    # 3. --- Review Trend Analysis ---
    review_trend_analysis_heading = document.add_heading('Review Trend Analysis', level=2)
    for run in review_trend_analysis_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    trend_pos_to_neg_heading = document.add_heading('Top 3 Shifts from Positive to Negative', level=3)
    for run in trend_pos_to_neg_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can analyze trends over time
    # For simplicity, we'll use trend analysis from the latest month
    if data:
        process_text(document, str(data[-1][1]['trend_pos_to_neg']) if data[-1][1]['trend_pos_to_neg'] else "No data available")
    else:
        document.add_paragraph("No data available")

    trend_neg_to_pos_heading = document.add_heading('Top 3 Shifts from Negative to Positive', level=3)
    for run in trend_neg_to_pos_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can analyze trends over time
    # For simplicity, we'll use trend analysis from the latest month
    if data:
        process_text(document, str(data[-1][1]['trend_neg_to_pos']) if data[-1][1]['trend_neg_to_pos'] else "No data available")
    else:
        document.add_paragraph("No data available")

    document.add_page_break()
    # 4. --- Competitor Analysis ---
    competitor_analysis_heading = document.add_heading('Competitor Analysis', level=2)
    for run in competitor_analysis_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    where_i_do_better_heading = document.add_heading('Where I Do Better', level=3)
    for run in where_i_do_better_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can analyze how competitive advantages change over time
    # For simplicity, we'll use competitor analysis from the latest month
    if data:
        process_text(document, str(data[-1][1]['where_i_do_better']) if data[-1][1]['where_i_do_better'] else "No data available")
    else:
        document.add_paragraph("No data available")

    where_competitor_do_better_heading = document.add_heading('Where Competitor Does Better', level=3)
    for run in where_competitor_do_better_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    # For multi-month analysis, you can analyze how competitive advantages change over time
    # For simplicity, we'll use competitor analysis from the latest month
    if data:
        process_text(document, str(data[-1][1]['where_competitor_do_better']) if data[-1][1]['where_competitor_do_better'] else "No data available")
    else:
        document.add_paragraph("No data available")

    # Set justify alignment for all paragraphs in the document
    for paragraph in document.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    document.save(output_filename)
    print(f"Word document '{output_filename}' created successfully.")


def open_word_file(word_file_path):
    if os.path.exists(word_file_path):
        subprocess.run(['open', word_file_path])
    else:
        logging.error(f"File '{word_file_path}' does not exist.")

# --- Main execution ---
# if __name__ == "__main__":
#     outlet_value = "South Plainfield"  # Replace with the desired outlet value
#     review_month_value = 2  # Replace with the desired review month value
#     num_months = 3 # setting it to three.

#     data = fetch_data_from_db(outlet_value, review_month_value, num_months=num_months)
#     path = f"A2B_{outlet_value}_{review_month_value}.docx"
#     if data is not None:
#        word =  create_word_document(outlet_value, review_month_value, data,
#                              output_filename=path)
#        open_word_file(path)
#     else:
#         print("Failed to fetch data. Check credentials and query.")